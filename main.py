import sys
print(sys.executable)

import settings as settings_manager
from tagVideo import *
from createMedia import renderPage3
from createClips import *
from styles import *
import base64
import io
import re
import copy
from dash import html, dcc, Input, Output, State
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


app.clientside_callback(
    """
    function(children) {
        var terminal = document.getElementById('terminal');
        if (terminal) {
            terminal.scrollTop = terminal.scrollHeight;
        }
        return window.dash_clientside.no_update;
    }
    """,
    Output("terminal", "style", allow_duplicate=True),
    Input("terminal", "children"),
    prevent_initial_call=True
)

app.clientside_callback(
    """
    function(n) {
        var panel = document.getElementById('settings-panel');
        if (panel.style.right === '0px') {
            panel.style.right = '-600px';
        } else {
            panel.style.right = '0px';
        }
        return window.dash_clientside.no_update;
    }
    """,
    Output("settings-panel", "style", allow_duplicate=True),
    Input("settings-btn", "n_clicks"),
    prevent_initial_call=True,
)

app.clientside_callback(
    """
    function(n) {
        var panel = document.getElementById('settings-panel');
        panel.style.right = '-600px';
        return window.dash_clientside.no_update;
    }
    """,
    Output("settings-panel", "style"),
    Input("save-settings-btn", "n_clicks"),
    prevent_initial_call=True,
)

app.clientside_callback(
    """
    function(page) {
        const videos = document.querySelectorAll('video');
        videos.forEach(v => {
            v.pause();
            v.src = '';
            v.load();
        });
        return window.dash_clientside.no_update;
    }
    """,
    Output("search-results-panel", "data-dummy"),
    Input("search-page-store", "data"),
    prevent_initial_call=True,
)

def renderPage1():
    return html.Div(
            style={
                "display": "flex",
                "flexDirection": "column",
                "alignItems": "center",
                "justifyContent": "center",
                "height": "100%",
                "gap": "20px",
            },
            children=[
                html.H1(
                    "ClipLocker",
                    style=headerStyle
                ),
                html.Div(
                    style={
                        "display": "flex",
                        "flexDirection": "column",
                        "alignItems": "center",
                        "gap": "20px",
                        "marginTop": "20px"
                    },
                    children=[
                        html.Button("Neue Clips taggen", id="btn-page2", style={**titleBtnStyle, "width": "400px", "height": "60px"}),
                        html.Button("Video-Clip Ordner erstellen", id="btn-page3", style={**titleBtnStyle, "width": "400px", "height": "60px"}),
                        html.Button("Clips aus Youtube-Video erstellen", id="btn-page4", style={**titleBtnStyle, "width": "400px", "height": "60px"}),
                    ]
                )
            ]
        )

def serveLayout():
    current_settings = settings_manager.load_settings()

    return html.Div(
        style={
            "display": "flex",
            "flexDirection": "column",
            "height": "100vh",
            # "backgroundColor": "#1e1e1e",
            "margin": "0",
            "padding": "0",
        },
        children=[

            html.Link(
                rel="stylesheet",
                href="https://fonts.googleapis.com/css2?family=Rajdhani:wght@500;700&display=swap"
            ),

            #global
            dcc.Interval(id="log-interval", interval=1000),
            dcc.Store(id="app-settings-store", data=current_settings, storage_type="memory"),
            dcc.Store(id="page-store", data="page1", storage_type="memory"),
            dcc.Store(id="epoch-store", data="medieval", storage_type="memory"),
            dcc.Store(id="session-store", storage_type="session"),

            #Tag Clips
            dcc.Store(id="clips-store", data=[], storage_type="memory"),
            dcc.Store(id="pathInput-store", data=None, storage_type="memory"),
            dcc.Store(id="currentClip-store", data=None, storage_type="memory"),
            dcc.Store(id="clip-settings-store", data={}, storage_type="memory"),
            dcc.Store(id="claude-descriptions-store", data=None, storage_type="memory"),

            #Ordner erstellen
            dcc.Store(id="sentence-store", data=[], storage_type="memory"),
            dcc.Store(id="clips-per-sentence-store",data=[], storage_type="memory"),
            dcc.Store(id="selectedClips-store",data=[], storage_type="memory"),
            dcc.Store(id="selectedSentence-store", data=None, storage_type="memory"),
            dcc.Store(id="documentName-store",data=None, storage_type="memory"),
            dcc.Store(id="active-right-tab", data="search", storage_type="memory"),
            dcc.Store(id="search-results-store", data=[], storage_type="memory"),
            dcc.Store(id="search-mode-store", data="and", storage_type="memory"),
            dcc.Store(id="search-page-store", data=0, storage_type="memory"),
            dcc.Store(id="placeholder-active-store", data=[], storage_type="memory"),
            dcc.Store(id="clip-usage-store", data={}, storage_type="memory"),
            dcc.Store(id="edit-tags-store", data=None, storage_type="memory"),
            dcc.Store(id="edit-tags-draft-store", data=None, storage_type="memory"),
            dcc.Store(id="edit-clip-id-store", data=None, storage_type="memory"),
            dcc.Download(id="clip-download"),


            #clips erstellen
            dcc.Store(id="p4-clips-store", data=[], storage_type="memory"),
            dcc.Store(id="p4-selected-store", data=None, storage_type="memory"),
            dcc.Store(id="p4-current-time-store", data=0, storage_type="memory"),
            dcc.Store(id="p4-split-time-store", data=None, storage_type="memory"),
            dcc.Store(id="p4-trim-start-time-store", data=None, storage_type="memory"),
            dcc.Store(id="p4-trim-end-time-store", data=None, storage_type="memory"),
            dcc.Store(id="p4-time-store", data=0, storage_type="memory"),

            #dummy
            html.Div(id="p4-split-time-label", style={"display": "none"}),
            html.Div(id="p4-bar-inner", style={"display": "none"}),
            html.Div(id="p4-pct-label", style={"display": "none"}),
            html.Div(id="p4-status-label", style={"display": "none"}),
            html.Div(id="p4-progress-area", style={"display": "none"}),
            html.Div(id="p4-input-area", style={"display": "none"}),
            html.Div(id="p4-strip", style={"display": "none"}),
            html.Div(id="p4-large-player", style={"display": "none"}),
            html.Div(id="p4-selected-label", style={"display": "none"}),
            html.Div(id="p4-trim-area", style={"display": "none"}),
            html.Div(id="p4-trim-start-label", style={"display": "none"}),
            html.Div(id="p4-trim-end-label", style={"display": "none"}),
            html.Div(id="p4-trim-slider", style={"display": "none"}),
            html.Img(id="p4-frame-start", style={"display": "none"}),
            html.Img(id="p4-frame-end", style={"display": "none"}),
            html.Div(id="p4-save-status", style={"display": "none"}),
            html.Div(id="p4-time-display", style={"display": "none"}),
            html.Div(id="main-container", style={"display": "none"}),
            html.Div(id="p4-prev-btn", style={"display": "none"}),
            html.Div(id="p4-next-btn", style={"display": "none"}),
            dcc.Input(id="p4-folder-input", style={"display": "none"}),
            html.Button(id="p4-folder-btn", style={"display": "none"}),
            html.Button(id="p4-save-btn", style={"display": "none"}),
            dcc.Input(id="p4-output-folder-input", style={"display": "none"}),
            html.Button(id="p4-trim-03-btn", style={"display": "none"}),
            html.Button(id="keyboard-enter-btn", style={"display": "none"}, n_clicks=0),
            html.Div(id="confirm-success-msg", style={"display": "none"}),
            html.Div(id="create-folder-success-msg", style={"display": "none"}),
            dcc.Store(id="descriptions-store", data=[], storage_type="memory"),



            html.Div(
                style={"display": "flex", "justifyContent": "flex-end", "padding": "10px"},
                children=[
                    html.Button(
                        "⚙",
                        id="settings-btn",
                        style=titleBtnStyle,
                    ),
                ]
            ),
            html.Div(
                id="page-content",
                style={"display": "flex", "flexDirection": "column", "flex": "1", "overflow": "hidden",
                       "minHeight": "0", }
            ),

            html.Div(
                id="terminal",
                style=terminalStyle
            ),
            html.Div(
                id="settings-panel",
                style={
                    "position": "fixed",
                    "top": "0",
                    "right": "-600px",
                    "width": "600px",
                    "height": "100vh",
                    "backgroundColor": "#2a2a2a",
                    "borderLeft": "1px solid #444",
                    "zIndex": "1000",
                    "transition": "right 0.3s ease",
                    "padding": "20px",
                    "boxSizing": "border-box",
                },
                children=[
                    #html.H2("Paths", style={"color": "#d4d4d4"}),
                    html.Label("Path:", style={"color": "#d4d4d4"}),
                    dcc.Input(
                        id={"type": "settings-input", "key": "path"},
                        type="text",
                        style={**baseStyleInputPath, "marginBottom": "10px"},
                        value=current_settings.get("path", ""),
                    ),
                    html.Label("Datenbank:", style={"color": "#d4d4d4"}),
                    dcc.Input(
                        id={"type": "settings-input", "key": "db"},
                        type="text",
                        style={**baseStyleInputPath, "marginBottom": "10px"},
                        value=current_settings.get("db", ""),
                    ),
                    html.Label("Fotosammlung:", style={"color": "#d4d4d4"}),
                    dcc.Input(
                        id={"type": "settings-input", "key": "fs"},
                        type="text",
                        style={**baseStyleInputPath, "marginBottom": "10px"},
                        value=current_settings.get("fs", ""),
                    ),
                    html.Label("Raw-Folder:", style={"color": "#d4d4d4"}),
                    dcc.Input(
                        id={"type": "settings-input", "key": "raw"},
                        type="text",
                        style={**baseStyleInputPath, "marginBottom": "10px"},
                        value=current_settings.get("raw", ""),
                    ),
                    html.Label("Db-Backups:", style={"color": "#d4d4d4"}),
                    dcc.Input(
                        id={"type": "settings-input", "key": "backup"},
                        type="text",
                        style={**baseStyleInputPath, "marginBottom": "10px"},
                        value=current_settings.get("backup", ""),
                    ),
                    html.Label("Temp-Folder:", style={"color": "#d4d4d4"}),
                    dcc.Input(
                        id={"type": "settings-input", "key": "temp"},
                        type="text",
                        style={**baseStyleInputPath, "marginBottom": "10px"},
                        value=current_settings.get("temp", ""),
                    ),
                    html.Button(
                        "Einstellungen speichern",
                        id="save-settings-btn",
                        style={**titleBtnStyle, "marginTop": "30px"},
                    ),
                ]
            ),
            html.Div(
                id ="word-wrapper",
                style={
                    "position": "fixed",
                    "bottom": "20px",
                    "left": "20px",
                    "zIndex": "999",
                    "height": "80px",
                    "width": "200px",
                },
                children=[
                    dcc.Upload(
                        id="word-upload-input",
                        children=html.Div(
                            ".docx ablegen oder auswählen",
                            style={
                                "color": "#888",
                                "fontSize": "12px",
                                "letterSpacing": "0.4px",
                            }
                        ),
                        accept=".docx",
                        multiple=False,
                        style={
                            "border": "1px dashed #555",
                            "borderRadius": "6px",
                            "padding": "10px 16px",
                            "cursor": "pointer",
                            "backgroundColor": "transparent",
                        },
                    ),
                    html.Div(
                        id="word-upload-status",
                        style={
                            "marginTop": "5px",
                            "fontSize": "11px",
                            "color": "#666",
                            "textAlign": "right",
                        }
                    ),
                    dcc.Download(id="word-download"),
                    dcc.Download(id="word-table-download"),
                ]
            )
        ]
    )


app.layout = serveLayout


@app.callback(
    Output("app-settings-store", "data"),
    Input("save-settings-btn", "n_clicks"),
    State({"type": "settings-input", "key": ALL}, "value"),
    State({"type": "settings-input", "key": ALL}, "id"),
    prevent_initial_call=True,
)
def save_settings_callback(n, values, ids):
    new_settings = {item["key"]: value for item, value in zip(ids, values)}
    saved_settings = settings_manager.save_settings(new_settings)
    return saved_settings

@app.callback(
    Output({"type": "settings-input", "key": ALL}, "value"),
    Input("app-settings-store", "data"),
    State({"type": "settings-input", "key": ALL}, "id"),
)
def sync_settings_inputs(appSettings, ids):
    return [appSettings.get(item["key"], "") for item in ids]


@app.callback(
    Output("page-store", "data"),
    Input("btn-page2", "n_clicks"),
    Input("btn-page3", "n_clicks"),
    Input("btn-page4", "n_clicks"),
    prevent_initial_call=True
)
def switchPage(p2, p3, p4):
    trigger = ctx.triggered_id
    if trigger == "btn-page2": return "page2_1"
    if trigger == "btn-page3": return "page3"
    if trigger == "btn-page4": return "page4"

    return "page1"

@app.callback(
    Output("page-content", "children"),
    Output("terminal", "style"),
    Output("settings-btn", "style"),
    Output("word-wrapper", "style"),
    Input("page-store", "data"),
    State("claude-descriptions-store", "data"),
    State("clips-store", "data"),
    State("app-settings-store", "data"),
)
def renderPage(page, claudeData, clips, appSettings):
    if page == "page2_1":
        return renderPage2_1(appSettings), {**terminalStyle, "display": "block"} , {**titleBtnStyle, "display": "none"}, {"display": "none"}
    if page == "page2_2":
        return renderPage2_2(claudeData, clips, appSettings), {**terminalStyle, "display": "block"}, {**titleBtnStyle, "display": "none"}, {"display": "none"}
    if page == "page3":
        return renderPage3(appSettings), {**terminalStyle, "display": "none"}, {**titleBtnStyle, "display": "none"}, {"display": "none"}
    if page == "page4":
        return renderPage4(appSettings), {**terminalStyle, "display": "none"}, {**titleBtnStyle, "display": "none"}, {"display": "none"}
    return renderPage1(), {"display": "none"}, {**titleBtnStyle, "display": "block"}, {"display": "block"}


def _set_no_spacing(p_el):
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_el.insert(0, pPr)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:before"), "0")


def _split_run_at_periods(paragraph):
    segments = []
    for run in paragraph.runs:
        rpr = run._r.find(qn("w:rPr"))
        rpr_xml = rpr.xml if rpr is not None else None
        parts = re.split(r"(\.\s*)", run.text)
        for part in parts:
            if not part:
                continue
            if re.fullmatch(r"\.\s*", part):
                if segments and not segments[-1][2]:
                    prev_text, prev_rpr, _ = segments[-1]
                    segments[-1] = (prev_text + ".", prev_rpr, True)
                else:
                    segments.append((".", rpr_xml, True))
            else:
                segments.append((part, rpr_xml, False))
    return segments


def format_docx_periods(file_bytes: bytes) -> bytes:
    doc = Document(io.BytesIO(file_bytes))

    for para in doc.paragraphs:
        if not para.text.strip():
            continue

        segments = _split_run_at_periods(para)
        if not any(is_break for _, _, is_break in segments):
            continue

        for run in para.runs:
            run.text = ""

        p = para._p
        parent = p.getparent()
        insert_idx = list(parent).index(p)

        current_p = p
        _set_no_spacing(current_p)

        for seg_text, rpr_xml, is_break in segments:
            if seg_text:
                r_el = OxmlElement("w:r")
                if rpr_xml:
                    from lxml import etree
                    r_el.append(etree.fromstring(rpr_xml))
                t_el = OxmlElement("w:t")
                t_el.text = seg_text
                t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                r_el.append(t_el)
                current_p.append(r_el)

            if is_break:
                new_p = copy.deepcopy(p)
                for r in new_p.findall(qn("w:r")):
                    new_p.remove(r)
                _set_no_spacing(new_p)
                insert_idx += 1
                parent.insert(insert_idx, new_p)
                current_p = new_p

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()

def _clone_paragraph_properties(src_para, dst_para):
    """Kopiert die Absatzeigenschaften (pPr) vom Quell- in den Zielparagraphen."""
    src_pPr = src_para._p.find(qn("w:pPr"))
    if src_pPr is not None:
        dst_pPr = copy.deepcopy(src_pPr)
        # vorhandenes pPr entfernen und neues einsetzen
        existing = dst_para._p.find(qn("w:pPr"))
        if existing is not None:
            dst_para._p.remove(existing)
        dst_para._p.insert(0, dst_pPr)




# ──────────────────────────────────────────────
#  Callbacks
# ──────────────────────────────────────────────
import json
import anthropic

@app.callback(
    Output("word-download", "data"),
    Output("word-table-download", "data"),   # ← NEU
    Output("word-upload-status", "children"),
    Output("word-upload-status", "style"),
    Input("word-upload-input", "contents"),
    State("word-upload-input", "filename"),
    prevent_initial_call=True,
)
def process_word_upload(contents, filename):
    base_status_style = {
        "marginTop": "10px", "fontSize": "12px",
        "textAlign": "center", "minHeight": "16px", "letterSpacing": "0.3px",
    }

    if contents is None:
        return None, None, "", {**base_status_style, "color": "#4a7a9b"}

    if not (filename or "").lower().endswith(".docx"):
        return None, None, "⚠ Nur .docx Dateien erlaubt", {**base_status_style, "color": "#e05252"}

    try:
        content_type, content_string = contents.split(",", 1)
        file_bytes = base64.b64decode(content_string)

        # ── 1. Docx formatieren (bestehend) ──────────────────────────
        formatted_bytes = format_docx_periods(file_bytes)
        stem = re.sub(r"\.docx$", "", filename, flags=re.IGNORECASE)
        out_docx_name = f"{stem}_formatiert.docx"

        # ── 2. Sätze extrahieren ──────────────────────────────────────
        doc = Document(io.BytesIO(formatted_bytes))
        sentences = [p.text for p in doc.paragraphs if p.text.strip()]

        # ── 3. Beschreibungen via Claude API generieren ───────────────
        client = anthropic.Anthropic(
            api_key="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")

        prompt = (
            "Du bekommst eine Liste von Sätzen aus einem Video-Skript. "
            "Erstelle für jeden Satz eine kurze englische Clip-Beschreibung (10-20 Wörter), "
            "die beschreibt welche Art von Videoclip inhaltlich gut zu diesem Satz passen würde. "
            "Antworte NUR mit einem JSON-Array, ohne Erklärungen oder Markdown:\n"
            '[{"sentence": "...", "description": "..."}, ...]\n\n'
            f"Sätze:\n{json.dumps(sentences, ensure_ascii=False)}"
        )

        print("Claude ackert...")

        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=8000,
            messages=[{"role": "user", "content": prompt}]
        )

        raw = message.content[0].text.strip()
        # JSON-Fences entfernen falls vorhanden
        raw = re.sub(r"^```json\s*|^```\s*|```$", "", raw, flags=re.MULTILINE).strip()
        table_data = json.loads(raw)

        # Fallback: falls Länge nicht stimmt
        if len(table_data) != len(sentences):
            table_data = [{"sentence": s, "description": ""} for s in sentences]

        out_json_name = f"{stem}_tabelle.json"
        json_bytes = json.dumps(table_data, ensure_ascii=False, indent=2).encode("utf-8")

        return (
            dcc.send_bytes(formatted_bytes, out_docx_name),
            dcc.send_bytes(json_bytes, out_json_name),
            f"✓ {out_docx_name} + {out_json_name}",
            {**base_status_style, "color": "#4caf82"},
        )

    except Exception as exc:
        return None, None, f"✗ Fehler: {exc}", {**base_status_style, "color": "#e05252"}


if __name__ == "__main__":
    app.run(debug=True, dev_tools_hot_reload=False)